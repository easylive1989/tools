import os
from docx import Document
from rich.console import Console
from rich.progress import Progress, SpinnerColumn, TextColumn, BarColumn, TaskProgressColumn
from src.services.gemini import GeminiClient

console = Console()

class DocxHandler:
    def __init__(self, gemini_client: GeminiClient, target_lang: str):
        self.client = gemini_client
        self.target_lang = target_lang

    def process(self, input_path: str):
        """
        Reads a DOCX file, translates paragraphs and table cells,
        and saves the result.
        """
        output_path = f"{os.path.splitext(input_path)[0]}_translated.docx"

        doc = Document(input_path)

        # Calculate total work (Paragraphs + Table Cells)
        total_steps = len(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    total_steps += len(cell.paragraphs)

        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            BarColumn(),
            TaskProgressColumn(),
            console=console,
            transient=True
        ) as progress:
            task = progress.add_task("[green]Translating...", total=total_steps)

            # Translate Paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    # To preserve runs (bold/italic) is complex in Phase 1.
                    # MVP Strategy: Translate the whole paragraph text and replace it.
                    original_text = para.text
                    translated_text = self.client.translate_text(original_text, self.target_lang)

                    # Clear existing content
                    para.clear()
                    # Add translated text
                    para.add_run(translated_text)
                progress.advance(task)

            # Translate Tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # Cells can contain paragraphs too
                        for para in cell.paragraphs:
                            if para.text.strip():
                                original_text = para.text
                                translated_text = self.client.translate_text(original_text, self.target_lang)
                                para.clear()
                                para.add_run(translated_text)
                            progress.advance(task)

        doc.save(output_path)
        return output_path
