import os
import unittest
from unittest.mock import MagicMock
import sys

# Ensure project root is in path
project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
if project_root not in sys.path:
    sys.path.insert(0, project_root)

from src.handlers.markdown import MarkdownHandler
from src.handlers.docx import DocxHandler
from src.services.gemini import GeminiClient

class TestHandlers(unittest.TestCase):
    def setUp(self):
        self.mock_client = MagicMock(spec=GeminiClient)
        self.mock_client.translate_text.side_effect = lambda text, lang: f"[TR:{lang}] {text}"
        self.target_lang = "Spanish"
        self.test_files = []

    def tearDown(self):
        # Cleanup test files
        for f in self.test_files:
            if os.path.exists(f):
                os.remove(f)

    def test_markdown_handler(self):
        handler = MarkdownHandler(self.mock_client, self.target_lang)
        input_md = "test_handlers_sample.md"
        self.test_files.append(input_md)
        output_md = "test_handlers_sample_translated.md"
        self.test_files.append(output_md)

        with open(input_md, 'w') as f:
            f.write("# Hello\n\n```\ncode\n```")

        output_path = handler.process(input_md)
        self.assertTrue(os.path.exists(output_path))

        with open(output_path, 'r') as f:
            content = f.read()
            # Code block should be preserved
            self.assertIn("```", content)
            # Text should be translated (mocked)
            self.assertIn("[TR:Spanish]", content)

    def test_docx_handler(self):
        handler = DocxHandler(self.mock_client, self.target_lang)
        input_docx = "test_handlers_sample.docx"
        self.test_files.append(input_docx)
        output_docx = "test_handlers_sample_translated.docx"
        self.test_files.append(output_docx)

        from docx import Document
        doc = Document()
        doc.add_paragraph("Hello World")
        doc.save(input_docx)

        output_path = handler.process(input_docx)
        self.assertTrue(os.path.exists(output_path))

        doc = Document(output_path)
        # Check if text is translated (mocked)
        full_text = " ".join([p.text for p in doc.paragraphs])
        self.assertIn("[TR:Spanish]", full_text)


class TestMarkdownBilingualHandler(unittest.TestCase):
    def setUp(self):
        self.mock_client = MagicMock(spec=GeminiClient)
        self.mock_client.translate_texts.side_effect = lambda texts, lang, **kwargs: [
            f"[TR:{lang}] {t}" for t in texts
        ]
        self.target_lang = "Traditional Chinese"
        self.test_files = []

    def tearDown(self):
        for f in self.test_files:
            if os.path.exists(f):
                os.remove(f)

    def _run(self, content: str) -> str:
        from src.handlers.markdown_bilingual import MarkdownBilingualHandler
        handler = MarkdownBilingualHandler(self.mock_client, self.target_lang)
        input_md = "test_bilingual_sample.md"
        output_md = "test_bilingual_sample_translated.md"
        self.test_files.extend([input_md, output_md])

        with open(input_md, 'w', encoding='utf-8') as f:
            f.write(content)

        handler.process(input_md)
        with open(output_md, 'r', encoding='utf-8') as f:
            return f.read()

    def test_plain_text_has_original_and_translation(self):
        result = self._run("Hello World")
        self.assertIn("Hello World", result)
        self.assertIn("[TR:Traditional Chinese]", result)

    def test_image_not_translated(self):
        result = self._run("![alt](image.png)")
        self.assertIn("![alt](image.png)", result)
        self.assertNotIn("[TR:", result)

    def test_link_only_not_translated(self):
        result = self._run("[click here](https://example.com)")
        self.assertIn("[click here](https://example.com)", result)
        self.assertNotIn("[TR:", result)

    def test_code_block_not_translated(self):
        result = self._run("```python\nprint('hello')\n```")
        self.assertIn("```python\nprint('hello')\n```", result)
        self.assertNotIn("[TR:", result)

    def test_mixed_content(self):
        content = "# Title\n\n![img](a.png)\n\nSome text here\n\n```\ncode\n```"
        result = self._run(content)
        # Title should be translated (original + translation)
        self.assertIn("# Title", result)
        self.assertIn("[TR:Traditional Chinese] # Title", result)
        # Image kept as-is
        self.assertIn("![img](a.png)", result)
        # Text translated
        self.assertIn("Some text here", result)
        self.assertIn("[TR:Traditional Chinese] Some text here", result)
        # Code block preserved
        self.assertIn("```\ncode\n```", result)

    def test_obsidian_frontmatter_preserved(self):
        content = (
            '---\n'
            'title: "A Motorcycle for the Mind"\n'
            'source: "https://nav.al/ai"\n'
            'tags:\n'
            '  - "clippings"\n'
            '---\n'
            '\n'
            '**Naval:** Podcast recording is so stilted.\n'
            '\n'
            '**Nivi:** And we all know brains run better.'
        )
        result = self._run(content)
        # Frontmatter preserved exactly
        self.assertIn('title: "A Motorcycle for the Mind"', result)
        self.assertIn('source: "https://nav.al/ai"', result)
        self.assertIn('  - "clippings"', result)
        # Body text: original preserved and translation added
        self.assertIn("**Naval:** Podcast recording is so stilted.", result)
        self.assertIn("[TR:Traditional Chinese]", result)


if __name__ == '__main__':
    unittest.main()
