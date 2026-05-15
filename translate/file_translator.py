# /// script
# requires-python = ">=3.10"
# dependencies = [
#     "python-docx",
#     "pdf2docx",
# ]
# ///
"""
Translate a .docx or .pdf file into Traditional Chinese using the local
`gemini` CLI. Output is written next to the source file as
`<stem>_translated.docx` (PDF inputs are converted to DOCX first).

Usage:
    uv run file_translator.py <path>
"""

import os
import re
import subprocess
import sys
from concurrent.futures import ThreadPoolExecutor, as_completed

from docx import Document
from pdf2docx import Converter


GEMINI_MODEL = "gemini-2.5-flash"
TARGET_LANG = "Traditional Chinese"
MAX_WORKERS = 5
ANSI_RE = re.compile(r"\x1b\[[0-9;]*[mGKHFABCDJKsuhl]|\r")


def log(msg: str) -> None:
    print(msg, file=sys.stderr, flush=True)


def gemini_translate(text: str) -> str:
    if not text or not text.strip():
        return text
    prompt = (
        f"Translate the following text into {TARGET_LANG}. "
        "Maintain the original tone and style. "
        "Do not add any explanations or extra text. "
        "Just provide the translation.\n\n"
        f"Text: {text}"
    )
    env = os.environ.copy()
    env["PATH"] = "/opt/homebrew/bin:/usr/local/bin:" + env.get("PATH", "")
    result = subprocess.run(
        ["gemini", "-m", GEMINI_MODEL, "-o", "text", prompt],
        capture_output=True,
        text=True,
        timeout=180,
        env=env,
    )
    if result.returncode != 0:
        err = ANSI_RE.sub("", result.stderr).strip()
        raise RuntimeError(f"gemini CLI failed: {err[:300]}")
    return ANSI_RE.sub("", result.stdout).strip()


def translate_many(texts: list[str]) -> list[str]:
    results: list[str] = [""] * len(texts)
    done = 0
    total = len(texts)
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as ex:
        futures = {ex.submit(gemini_translate, t): i for i, t in enumerate(texts)}
        for fut in as_completed(futures):
            idx = futures[fut]
            results[idx] = fut.result()
            done += 1
            log(f"progress {done}/{total}")
    return results


def translate_docx(input_path: str) -> str:
    output_path = f"{os.path.splitext(input_path)[0]}_translated.docx"
    doc = Document(input_path)

    para_refs = []
    texts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            para_refs.append(para)
            texts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        para_refs.append(para)
                        texts.append(para.text)

    log(f"translating {len(texts)} paragraphs")
    translated = translate_many(texts)
    for para, new_text in zip(para_refs, translated):
        para.clear()
        para.add_run(new_text)
    doc.save(output_path)
    return output_path


def translate_pdf(input_path: str) -> str:
    temp_docx = f"{os.path.splitext(input_path)[0]}.docx"
    log("converting PDF to DOCX")
    cv = Converter(input_path)
    try:
        cv.convert(temp_docx)
    finally:
        cv.close()
    try:
        return translate_docx(temp_docx)
    finally:
        if os.path.exists(temp_docx):
            try:
                os.remove(temp_docx)
            except OSError:
                pass


def main() -> int:
    if len(sys.argv) != 2:
        log("usage: file_translator.py <path>")
        return 2
    path = sys.argv[1]
    if not os.path.isfile(path):
        log(f"file not found: {path}")
        return 1
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        out = translate_docx(path)
    elif ext == ".pdf":
        out = translate_pdf(path)
    else:
        log(f"unsupported format: {ext}")
        return 1
    print(out)
    return 0


if __name__ == "__main__":
    sys.exit(main())
